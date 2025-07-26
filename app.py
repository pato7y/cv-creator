from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
import io
import re
from datetime import datetime
import tempfile

app = Flask(__name__)

class WebCVCreator:
    def __init__(self):
        self.doc = Document()
        self.setup_document_formatting()
        self.setup_custom_styles()
        
    def setup_document_formatting(self):
        """Configure document margins and layout"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
    
    def setup_custom_styles(self):
        """Create professional custom styles"""
        styles = self.doc.styles
        
        # Main heading style (Name)
        if 'CV Name' not in [s.name for s in styles]:
            name_style = styles.add_style('CV Name', WD_STYLE_TYPE.PARAGRAPH)
            name_font = name_style.font
            name_font.name = 'Georgia'
            name_font.size = Pt(24)
            name_font.bold = True
            name_font.color.rgb = RGBColor(0, 32, 96)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_style.paragraph_format.space_after = Pt(6)
        
        # Section heading style
        if 'CV Section' not in [s.name for s in styles]:
            section_style = styles.add_style('CV Section', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Georgia'
            section_font.size = Pt(14)
            section_font.bold = True
            section_font.color.rgb = RGBColor(0, 32, 96)
            section_font.all_caps = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        if 'CV Contact' not in [s.name for s in styles]:
            contact_style = styles.add_style('CV Contact', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Calibri'
            contact_font.size = Pt(11)
            contact_font.color.rgb = RGBColor(64, 64, 64)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Position style
        if 'CV Position' not in [s.name for s in styles]:
            position_style = styles.add_style('CV Position', WD_STYLE_TYPE.PARAGRAPH)
            position_font = position_style.font
            position_font.name = 'Calibri'
            position_font.size = Pt(12)
            position_font.bold = True
            position_style.paragraph_format.space_after = Pt(3)
        
        # Company style
        if 'CV Company' not in [s.name for s in styles]:
            company_style = styles.add_style('CV Company', WD_STYLE_TYPE.PARAGRAPH)
            company_font = company_style.font
            company_font.name = 'Calibri'
            company_font.size = Pt(11)
            company_font.italic = True
            company_font.color.rgb = RGBColor(96, 96, 96)
            company_style.paragraph_format.space_after = Pt(6)
        
        # Body style
        if 'CV Body' not in [s.name for s in styles]:
            body_style = styles.add_style('CV Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
    
    def add_section_divider(self):
        """Add a subtle horizontal line as section divider"""
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pBdr.set(qn('w:bottom'), 'single')
        pBdr.set(qn('w:sz'), '4')
        pBdr.set(qn('w:space'), '1')
        pBdr.set(qn('w:color'), 'CCCCCC')
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(12)
    
    def create_cv_from_data(self, data):
        """Create CV from web form data"""
        
        # Header: Name and Title
        full_name = f"{data['firstName']} {data['lastName']}"
        name_para = self.doc.add_paragraph(full_name, style='CV Name')
        
        if data.get('professionalTitle'):
            title_para = self.doc.add_paragraph(data['professionalTitle'])
            title_para.style = self.doc.styles['CV Contact']
            title_para.runs[0].font.size = Pt(14)
            title_para.runs[0].font.italic = True
        
        # Contact Information
        contact_info = []
        contact_info.append(data['email'])
        contact_info.append(data['phone'])
        
        location_parts = [data['city']]
        if data.get('state'):
            location_parts.append(data['state'])
        if data.get('country'):
            location_parts.append(data['country'])
        contact_info.append(', '.join(location_parts))
        
        if data.get('linkedin'):
            contact_info.append(f"LinkedIn: {data['linkedin']}")
        if data.get('github'):
            contact_info.append(f"GitHub: {data['github']}")
        if data.get('portfolio'):
            contact_info.append(f"Portfolio: {data['portfolio']}")
        
        self.doc.add_paragraph(' | '.join(contact_info), style='CV Contact')
        self.add_section_divider()
        
        # Professional Summary
        if data.get('summary'):
            self.doc.add_paragraph('PROFESSIONAL SUMMARY', style='CV Section')
            self.doc.add_paragraph(data['summary'], style='CV Body')
            self.add_section_divider()
        
        # Work Experience
        if data.get('experiences'):
            self.doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='CV Section')
            
            for exp in data['experiences']:
                # Job title
                job_para = self.doc.add_paragraph(style='CV Position')
                job_para.add_run(exp['jobTitle']).bold = True
                if exp.get('employmentType', 'Full-time') != 'Full-time':
                    job_para.add_run(f" ({exp['employmentType']})")
                
                # Company and dates
                company_info = f"{exp['company']} | {exp['location']} | {exp['startDate']} - {exp['endDate']}"
                self.doc.add_paragraph(company_info, style='CV Company')
                
                # Responsibilities
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        if resp.strip():
                            bullet_para = self.doc.add_paragraph(f"• {resp}", style='CV Body')
                            bullet_para.paragraph_format.left_indent = Inches(0.25)
                
                self.doc.add_paragraph()
            
            self.add_section_divider()
        
        # Education
        if data.get('education'):
            self.doc.add_paragraph('EDUCATION', style='CV Section')
            
            for edu in data['education']:
                # Degree
                degree_para = self.doc.add_paragraph(style='CV Position')
                degree_para.add_run(f"{edu['degree']} in {edu['fieldOfStudy']}").bold = True
                
                # Institution
                inst_info = f"{edu['institution']}, {edu['location']}"
                if edu.get('graduationDate'):
                    inst_info += f" | Graduated: {edu['graduationDate']}"
                
                self.doc.add_paragraph(inst_info, style='CV Company')
                
                # GPA and Honors
                additional_info = []
                if edu.get('gpa'):
                    additional_info.append(f"GPA: {edu['gpa']}")
                if edu.get('honors'):
                    additional_info.append(f"Honors: {edu['honors']}")
                
                if additional_info:
                    self.doc.add_paragraph(' | '.join(additional_info), style='CV Body')
                
                self.doc.add_paragraph()
            
            self.add_section_divider()
        
        # Skills
        if data.get('skills'):
            self.doc.add_paragraph('CORE COMPETENCIES', style='CV Section')
            for skill_category in data['skills']:
                if skill_category.get('category') and skill_category.get('skills'):
                    skill_para = self.doc.add_paragraph(style='CV Body')
                    skill_para.add_run(f"{skill_category['category']}: ").bold = True
                    skill_para.add_run(' | '.join(skill_category['skills']))
            
            self.add_section_divider()
        
        # Projects
        if data.get('projects'):
            self.doc.add_paragraph('KEY PROJECTS', style='CV Section')
            
            for project in data['projects']:
                # Project name
                project_para = self.doc.add_paragraph(style='CV Position')
                project_para.add_run(project['name']).bold = True
                if project.get('type'):
                    project_para.add_run(f" ({project['type']} Project)")
                
                # Duration and URL
                project_info = []
                if project.get('duration'):
                    project_info.append(project['duration'])
                if project.get('url'):
                    project_info.append(project['url'])
                
                if project_info:
                    self.doc.add_paragraph(' | '.join(project_info), style='CV Company')
                
                # Description
                if project.get('description'):
                    self.doc.add_paragraph(project['description'], style='CV Body')
                
                # Technologies
                if project.get('technologies'):
                    tech_para = self.doc.add_paragraph(style='CV Body')
                    tech_para.add_run("Technologies: ").bold = True
                    tech_para.add_run(', '.join(project['technologies']))
                
                self.doc.add_paragraph()
            
            self.add_section_divider()
    
    def save_to_memory(self):
        """Save document to memory and return file-like object"""
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/create-cv', methods=['POST'])
def create_cv():
    try:
        data = request.json
        
        # Create CV
        cv_creator = WebCVCreator()
        cv_creator.create_cv_from_data(data)
        
        # Save to memory
        file_stream = cv_creator.save_to_memory()
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(file_stream.getvalue())
            temp_filename = tmp_file.name
        
        # Generate download filename
        full_name = f"{data['firstName']}_{data['lastName']}"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        download_filename = f"CV_{full_name}_{timestamp}.docx"
        
        return send_file(
            temp_filename,
            as_attachment=True,
            download_name=download_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up temporary file
        try:
            if 'temp_filename' in locals():
                os.unlink(temp_filename)
        except:
            pass

@app.route('/health')
def health_check():
    return jsonify({'status': 'healthy'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV') != 'production'
    app.run(host='0.0.0.0', port=port, debug=debug_mode)